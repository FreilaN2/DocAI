"""
core/document_builder.py
========================
Módulo de construcción de documentos Word con formato APA.
Extraído de main.py para mejorar la separación de responsabilidades
y reducir el tamaño del archivo principal.

Contiene:
- Constantes APA (fuentes, normas, márgenes)
- Helpers de texto (normalización, formato de títulos)
- Funciones de construcción DOCX (estilos, encabezados, marca de agua)
"""

import os
import re
import shutil
import subprocess
from functools import lru_cache
from typing import Optional

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING


# ═══════════════════════════════════════════════════════════
# CONSTANTES APA PRE-CALCULADAS (evita recalcular en cada request)
# ═══════════════════════════════════════════════════════════

LETTER_PAGE = {"width": Inches(8.5), "height": Inches(11.0)}

FUENTES_APA = {
    "Times New Roman": {"tamano": 12, "familia": "serif"},
    "Georgia":         {"tamano": 11, "familia": "serif"},
    "Computer Modern": {"tamano": 10, "familia": "serif"},
    "Calibri":         {"tamano": 11, "familia": "sans-serif"},
    "Arial":           {"tamano": 11, "familia": "sans-serif"},
    "Lucida Sans Unicode": {"tamano": 10, "familia": "sans-serif"},
}

DEFAULT_APA_FONT = "Times New Roman"

# Mapa pre-normalizado (lowercase → nombre original) para validar fuentes en O(1)
# sin iterar ni hacer lower() por cada entrada en cada request.
_FUENTES_APA_LOWER: dict[str, str] = {k.lower(): k for k in FUENTES_APA}

NORMAS_APA = {
    "6ta": {
        "fuente": DEFAULT_APA_FONT,
        "tamano": FUENTES_APA[DEFAULT_APA_FONT]["tamano"],
        "interlineado": 2.0,
        "margen": 1.0,
        "page_width": LETTER_PAGE["width"],
        "page_height": LETTER_PAGE["height"],
        "sangria_primera_linea": 0.5,
        "sangria_francesa": 0.5,
        "títulos": {
            "N1": {"bold": True,  "italic": False, "align": "center"},
            "N2": {"bold": True,  "italic": False, "align": "left"},
            "N3": {"bold": True,  "italic": False, "align": "left"},
            "N4": {"bold": True,  "italic": False, "align": "indent"},
            "N5": {"bold": True,  "italic": True,  "align": "indent"},
        },
    },
    "7ma": {
        "fuente": DEFAULT_APA_FONT,
        "tamano": FUENTES_APA[DEFAULT_APA_FONT]["tamano"],
        "interlineado": 2.0,
        "margen": 1.0,
        "page_width": LETTER_PAGE["width"],
        "page_height": LETTER_PAGE["height"],
        "sangria_primera_linea": 0.5,
        "sangria_francesa": 0.5,
        "títulos": {
            "N1": {"bold": True,  "italic": False, "align": "center"},
            "N2": {"bold": True,  "italic": False, "align": "left"},
            "N3": {"bold": True,  "italic": True,  "align": "left"},
            "N4": {"bold": True,  "italic": False, "align": "indent"},
            "N5": {"bold": True,  "italic": True,  "align": "indent"},
        },
    },
}

# Conjunto de palabras no-mayúsculas como constante de módulo (no se recrea en cada llamada)
_PALABRAS_NO_MAYUSCULAS = frozenset({
    "a", "ante", "bajo", "con", "contra", "de", "del", "desde",
    "durante", "e", "el", "la", "las", "los", "para", "por",
    "sin", "sobre", "y", "o", "u", "en", "al", "aun",
})

# Expresiones regulares compiladas una sola vez al cargar el módulo.
# Compilar dentro de las funciones crea un nuevo objeto regex en cada llamada.
_RE_NORMALIZAR_CAT_CHARS  = re.compile(r"[^A-Z0-9_]+")
_RE_NORMALIZAR_CAT_SPACES = re.compile(r"\s+")
_RE_NORMALIZAR_TEXTO_CHARS = re.compile(r"[^a-z0-9 ]+")
_RE_AUTOR_MATCH = re.compile(r'^\s*([A-ZÁÉÍÓÚÑÜÇ][\wÁÉÍÓÚÑÜÇ\'-]+)', re.IGNORECASE)

# Tabla de transliteración para acentos (creada una sola vez)
_TRANS_ACENTOS = str.maketrans("áéíóúüñç", "aeiouunc")

# Conjuntos pre-calculados de encabezados de referencias válidos
_ENCABEZADOS_REFERENCIAS = frozenset({
    "referencias",
    "referencia",
    "referencias bibliograficas",
    "referencia bibliografica",
    "references bibliograficas",
    "bibliografia",
    "bibliografias",
    "bibliograficas",
    "bibliografica",
})

_CONTINUACIONES_REFERENCIAS = frozenset({
    "bibliografia",
    "bibliografias",
    "bibliograficas",
    "bibliografica",
    "bibliografico",
})


# ═══════════════════════════════════════════════════════════
# HELPERS APA — TEXTO Y FORMATO
# ═══════════════════════════════════════════════════════════

def validar_fuente_apa(nombre_fuente: str) -> str:
    """Valida y normaliza el nombre de fuente APA. Lookup O(1) con dict pre-normalizado."""
    if not nombre_fuente:
        return DEFAULT_APA_FONT
    return _FUENTES_APA_LOWER.get(nombre_fuente.strip().lower(), DEFAULT_APA_FONT)


def formatear_titulo_apa(texto: str) -> str:
    """Aplica Title Case APA (respeta palabras menores como 'de', 'el', 'y')."""
    if not texto:
        return texto.strip()
    # Preservar mayúsculas sostenidas intencionales (ej. portadas institucionales o CAPÍTULOS)
    if texto.isupper():
        return texto.strip()
        
    partes = texto.strip().split()
    resultado = []
    for idx, palabra in enumerate(partes):
        lower = palabra.strip().lower()
        if idx == 0 or lower not in _PALABRAS_NO_MAYUSCULAS:
            resultado.append(lower.capitalize())
        else:
            resultado.append(lower)
    return " ".join(resultado)


def formatear_titulo_apa_sentence_case(texto: str) -> str:
    """Aplica Sentence case APA (solo primera letra mayúscula)."""
    if not texto:
        return texto.strip()
    # Preservar mayúsculas sostenidas intencionales
    if texto.isupper():
        return texto.strip()
        
    t = texto.strip().lower()
    return (t[0].upper() + t[1:]) if len(t) > 1 else t.upper()


def formatear_titulo_por_nivel(texto: str, nivel: str, edicion: str) -> str:
    """Selecciona el formato de título correcto según nivel y edición APA."""
    if edicion == "6ta" and nivel in {"N3", "N4", "N5"}:
        return formatear_titulo_apa_sentence_case(texto)
    return formatear_titulo_apa(texto)


def _normalizar_categoria(categoria: str) -> str:
    """Normaliza una categoría de párrafo al formato canónico interno."""
    if not categoria:
        return "PARRAFO_NORMAL"

    cat = _RE_NORMALIZAR_CAT_CHARS.sub("", categoria.strip().upper().replace(" ", "_"))

    if cat.startswith("TITULO") and "_N" not in cat:
        cat = "TITULO_N1"

    if cat.startswith("REFERENCIA"):
        return "REFERENCIA"
    if cat in {"CITA_LARGA", "BLOQUE_CITA"}:
        return cat
    if cat in {"TITULO_N1", "TITULO_N2", "TITULO_N3", "TITULO_N4", "TITULO_N5", "PARRAFO_NORMAL"}:
        return cat

    return "PARRAFO_NORMAL"


def _normalizar_texto_para_encabezado(texto: str) -> str:
    """Normaliza texto para comparación de encabezados (sin acentos, lowercase, sin especiales)."""
    t = texto.strip().lower().translate(_TRANS_ACENTOS)
    return _RE_NORMALIZAR_TEXTO_CHARS.sub("", t).strip()


def _es_encabezado_referencias(texto: str) -> bool:
    """Determina si un párrafo es el encabezado de la sección de referencias."""
    base = _normalizar_texto_para_encabezado(texto)
    if base in _ENCABEZADOS_REFERENCIAS:
        return True
    return ("referencia" in base and "bibliograf" in base) or base.startswith("referencia")


def _es_continuacion_encabezado_referencias(texto: str) -> bool:
    """Determina si un párrafo es la continuación del encabezado de referencias."""
    base = _normalizar_texto_para_encabezado(texto)
    return base in _CONTINUACIONES_REFERENCIAS or base.startswith("bibliograf")


def _ordenar_referencia_por_autor(texto: str) -> str:
    """Extrae la clave de ordenamiento de una referencia (apellido del primer autor)."""
    m = _RE_AUTOR_MATCH.match(texto.strip())
    if m:
        return m.group(1).lower()
    return _RE_NORMALIZAR_TEXTO_CHARS.sub("", texto.strip().lower())


# ═══════════════════════════════════════════════════════════
# ESTILOS DE PÁRRAFO APA
# ═══════════════════════════════════════════════════════════

def configurar_parrafo_estilo(paragraph, categoria: str, reglas: dict, body_text: str = None) -> None:
    """
    Aplica el estilo APA correspondiente a un párrafo de documento Word.

    Args:
        paragraph: Objeto párrafo de python-docx.
        categoria: Categoría del párrafo (TITULO_N1, REFERENCIA, PARRAFO_NORMAL, etc.).
        reglas: Dict con las reglas APA activas (fuente, tamaño, sangrías, etc.).
        body_text: Texto del cuerpo que se fusiona con títulos N3-N5 en 6ta edición.
    """
    categoria = _normalizar_categoria(categoria)
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(0)
    pf.keep_together     = False
    pf.first_line_indent = Inches(0)
    pf.left_indent       = Inches(0)
    paragraph.alignment  = WD_ALIGN_PARAGRAPH.JUSTIFY

    if categoria.startswith("TITULO"):
        nivel   = categoria.split("_")[-1] if "_" in categoria else "N1"
        edicion = reglas.get("edicion", "7ma")
        paragraph.text       = formatear_titulo_por_nivel(paragraph.text, nivel, edicion)
        pf.space_before      = Pt(12)
        pf.space_after       = Pt(0)
        pf.keep_with_next    = True
        pf.first_line_indent = Inches(0)
        pf.left_indent       = Inches(0)

        heading_map = {"N1": "Heading 1", "N2": "Heading 2"}
        if edicion == "6ta" and nivel in {"N3", "N4", "N5"}:
            paragraph.style      = "Normal"
            pf.left_indent       = Inches(reglas["sangria_primera_linea"])
            pf.first_line_indent = Inches(0)
        else:
            paragraph.style = heading_map.get(nivel, "Heading 1")
            if nivel in {"N3", "N4", "N5"}:
                pf.left_indent = Inches(reglas["sangria_primera_linea"])

        bold_italic_align = {
            "N1": (True,  False, WD_ALIGN_PARAGRAPH.CENTER),
            "N2": (True,  False, WD_ALIGN_PARAGRAPH.LEFT),
            "N3": (True,  False, WD_ALIGN_PARAGRAPH.LEFT),
            "N4": (True,  True,  WD_ALIGN_PARAGRAPH.LEFT),
            "N5": (False, True,  WD_ALIGN_PARAGRAPH.LEFT),
        }
        bold, italic, align = bold_italic_align.get(nivel, (True, False, WD_ALIGN_PARAGRAPH.LEFT))
        paragraph.alignment = align

        # Títulos N3-N5 en 6ta edición: título + punto + texto en el mismo párrafo
        if body_text and edicion == "6ta" and nivel in {"N3", "N4", "N5"}:
            heading_text = formatear_titulo_por_nivel(paragraph.text, nivel, edicion)
            if not heading_text.strip().endswith('.'):
                heading_text = heading_text.rstrip() + '.'
            paragraph.text = ""
            hr = paragraph.add_run(heading_text)
            hr.bold, hr.italic, hr.font.name, hr.font.size = (
                bold, italic, reglas["fuente"], Pt(reglas["tamano"])
            )
            paragraph.add_run(" ")
            br = paragraph.add_run(body_text.strip())
            br.bold, br.italic, br.font.name, br.font.size = (
                False, False, reglas["fuente"], Pt(reglas["tamano"])
            )
            return

        if edicion == "6ta" and nivel in {"N3", "N4", "N5"} and not paragraph.text.strip().endswith('.'):
            paragraph.text = paragraph.text.rstrip() + '.'

        from docx.shared import RGBColor
        for run in paragraph.runs:
            run.bold, run.italic = bold, italic
            run.font.name, run.font.size = reglas["fuente"], Pt(reglas["tamano"])
            run.font.color.rgb = RGBColor(0, 0, 0)
        return

    if categoria in {"CITA_LARGA", "BLOQUE_CITA"}:
        paragraph.text      = paragraph.text.strip().strip('""„»')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.left_indent      = Inches(reglas["sangria_primera_linea"])
        pf.first_line_indent = Inches(0)
        for run in paragraph.runs:
            run.font.name, run.font.size = reglas["fuente"], Pt(reglas["tamano"])
        return

    if categoria == "REFERENCIA":
        paragraph.alignment  = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Inches(-reglas["sangria_francesa"])
        pf.left_indent       = Inches(reglas["sangria_francesa"])
        pf.keep_together     = False
        pf.space_before      = Pt(0)
        pf.space_after       = Pt(0)
        for run in paragraph.runs:
            run.font.name, run.font.size = reglas["fuente"], Pt(reglas["tamano"])
        return

    # PARRAFO_NORMAL (default)
    pf.first_line_indent = Inches(reglas["sangria_primera_linea"])
    paragraph.alignment  = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        run.font.name, run.font.size = reglas["fuente"], Pt(reglas["tamano"])


# ═══════════════════════════════════════════════════════════
# UTILIDADES DE DOCUMENTO WORD
# ═══════════════════════════════════════════════════════════

def añadir_marca_de_agua(doc) -> None:
    """Inserta una marca de agua en el pie de página (plan Free)."""
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Generado con DocAI Free — Formateador Automático APA")
        run.font.size = Pt(10)
        run.font.name = "Arial"


def _insertar_tabla_de_contenidos(doc) -> None:
    """Inserta un campo TOC de Word para la tabla de contenidos automática."""
    paragraph = doc.add_paragraph()
    fld_simple = OxmlElement('w:fldSimple')
    fld_simple.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    fld_simple.set(qn('w:dirty'), 'true')
    paragraph._p.append(fld_simple)


def _configurar_encabezado_paginas(doc) -> None:
    """Configura el encabezado de páginas con número de página centrado."""
    for i, section in enumerate(doc.sections):
        section.header.is_linked_to_previous = False
        
        # Ocultar el número en la primera página (portada)
        if i == 0:
            section.different_first_page_header_footer = True
            
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run()
        
        # Estructura compleja de campo PAGE para que LibreOffice (y Word) rendericen un número por defecto (fallback '1')
        # hasta que se actualicen los campos, garantizando que no salga vacío en el PDF final.
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        r_fallback = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = "1"
        r_fallback.append(t)
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(r_fallback)
        run._r.append(fldChar3)


def _force_update_fields(doc) -> None:
    """Fuerza la actualización de campos (PAGE, TOC) al abrir en Word."""
    settings = getattr(doc, 'settings', None)
    if not settings:
        return
    element = getattr(settings, 'element', None)
    if element is None:
        return
    update = OxmlElement('w:updateFields')
    update.set(qn('w:val'), 'true')
    element.append(update)


def copiar_parrafo_xml(doc_original, doc_nuevo, idx: int):
    """
    Copia un único párrafo (por índice) desde el documento original al nuevo,
    resolviendo y reescribiendo las referencias a imágenes (rIds).
    """
    import copy

    _NS_R       = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    _REL_IMAGE  = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'
    _BLIP_TAG   = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'

    if not doc_original or idx >= len(doc_original.paragraphs):
        return

    orig_part  = doc_original.part
    nuevo_part = doc_nuevo.part
    orig_p = doc_original.paragraphs[idx]._p
    xml_copy = copy.deepcopy(orig_p)

    rId_map: dict[str, str] = {}

    for blip in xml_copy.iter(_BLIP_TAG):
        r_embed = blip.get(f'{{{_NS_R}}}embed')
        if not r_embed:
            continue
        if r_embed in rId_map:
            blip.set(f'{{{_NS_R}}}embed', rId_map[r_embed])
            continue
        try:
            img_part = orig_part.related_parts.get(r_embed)
            if img_part is None:
                continue
            nuevo_rId = nuevo_part.relate_to(img_part, _REL_IMAGE)
        except Exception:
            try:
                from docx.opc.part import Part
                from docx.opc.packuri import PackURI
                img_part2 = orig_part.related_parts.get(r_embed)
                if img_part2 is None:
                    continue
                new_p = Part(
                    PackURI(img_part2.partname),
                    img_part2.content_type,
                    img_part2._blob,
                    nuevo_part.package,
                )
                nuevo_rId = nuevo_part.relate_to(new_p, _REL_IMAGE)
            except Exception as exc:
                logger.warning(f"Imagen {r_embed} no copiada: {exc}")
                continue
        rId_map[r_embed] = nuevo_rId
        blip.set(f'{{{_NS_R}}}embed', nuevo_rId)

    # Limpiar propiedades originales del párrafo (ej. sangrías manuales o alineación heredada)
    # para que la nueva alineación pueda aplicarse limpiamente.
    if xml_copy.pPr is not None:
        xml_copy.remove(xml_copy.pPr)

    _WNS_SECTPR = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr'
    sect_pr = doc_nuevo.element.body.find(_WNS_SECTPR)
    if sect_pr is not None:
        sect_pr.addprevious(xml_copy)
    else:
        doc_nuevo.element.body.append(xml_copy)
        
    from docx.text.paragraph import Paragraph
    return Paragraph(xml_copy, doc_nuevo._body)


def _get_soffice_path() -> Optional[str]:
    """Detecta la ruta del ejecutable de LibreOffice en el sistema."""
    path_in_env = shutil.which("soffice")
    if path_in_env:
        return path_in_env
    for p in [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]:
        if os.path.exists(p):
            return p
    return None


# Resultado cacheado para no llamar shutil.which() en cada conversión PDF
_get_soffice_path_cached = lru_cache(maxsize=1)(_get_soffice_path)


# ═══════════════════════════════════════════════════════════
# COPIA DE PORTADA DESDE DOCUMENTO ORIGINAL (preserva imágenes)
# ═══════════════════════════════════════════════════════════

def _detectar_n_portada(parrafos) -> int:
    """
    Detecta cuántos párrafos iniciales corresponden a la portada.
    La portada termina justo antes del primer TITULO_N1 que corresponda
    al cuerpo real del documento (capítulo, el problema, resumen, etc.).
    
    Retorna el índice (0-based) del primer párrafo que NO es portada.
    Si no se detecta corte, retorna 0 (sin portada detectada).
    """
    # Palabras clave que indican inicio del cuerpo académico
    _INICIO_CUERPO = frozenset({
        "capitulo", "capítulo", "resumen", "abstract",
        "introduccion", "introducción", "el problema",
        "planteamiento", "agradecimientos", "dedicatoria",
        "indice", "índice", "referencias", "bibliograf",
    })

    for idx, p in enumerate(parrafos):
        cat = p.get("categoria", "") if isinstance(p, dict) else getattr(p, "categoria", "")
        txt = (p.get("texto", "") if isinstance(p, dict) else getattr(p, "texto", "")).strip().lower()

        is_title = cat.startswith("TITULO")
        is_short_normal = cat == "PARRAFO_NORMAL" and len(txt) < 100

        if is_title or is_short_normal:
            if any(txt.startswith(kw) for kw in _INICIO_CUERPO):
                return idx

        # Seguridad: si llegamos al párrafo 100 sin corte, no hay portada identificable
        if idx >= 100:
            return 0

    return 0

def copiar_parrafo_xml(doc_original, doc_nuevo, idx: int):
    import copy
    _NS_R       = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    _REL_IMAGE  = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'
    _BLIP_TAG   = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'

    if not doc_original or idx >= len(doc_original.paragraphs):
        return None

    orig_part  = doc_original.part
    nuevo_part = doc_nuevo.part
    orig_p = doc_original.paragraphs[idx]._p
    xml_copy = copy.deepcopy(orig_p)

    rId_map = {}

    for blip in xml_copy.iter(_BLIP_TAG):
        r_embed = blip.get(f'{{{_NS_R}}}embed')
        if not r_embed:
            continue
        if r_embed in rId_map:
            blip.set(f'{{{_NS_R}}}embed', rId_map[r_embed])
            continue
        try:
            img_part = orig_part.related_parts.get(r_embed)
            if img_part is None:
                continue
            nuevo_rId = nuevo_part.relate_to(img_part, _REL_IMAGE)
        except Exception:
            try:
                from docx.opc.part import Part
                from docx.opc.packuri import PackURI
                img_part2 = orig_part.related_parts.get(r_embed)
                if img_part2 is None:
                    continue
                new_p = Part(PackURI(img_part2.partname), img_part2.content_type, img_part2._blob, nuevo_part.package)
                nuevo_rId = nuevo_part.relate_to(new_p, _REL_IMAGE)
            except Exception:
                continue
        rId_map[r_embed] = nuevo_rId
        blip.set(f'{{{_NS_R}}}embed', nuevo_rId)

    if xml_copy.pPr is not None:
        xml_copy.remove(xml_copy.pPr)

    _WNS_SECTPR = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr'
    sect_pr = doc_nuevo.element.body.find(_WNS_SECTPR)
    if sect_pr is not None:
        sect_pr.addprevious(xml_copy)
    else:
        doc_nuevo.element.body.append(xml_copy)
        
    from docx.text.paragraph import Paragraph
    return Paragraph(xml_copy, doc_nuevo._body)


def copiar_elemento_xml(doc_original, doc_nuevo, element):
    """
    Copia cualquier elemento XML (párrafo, tabla, forma) conservando EXACTAMENTE
    su formato original (alineaciones, estilos, bordes).
    """
    import copy
    _NS_R       = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    _REL_IMAGE  = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'
    _BLIP_TAG   = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'

    if not doc_original or element is None:
        return None

    orig_part  = doc_original.part
    nuevo_part = doc_nuevo.part
    xml_copy = copy.deepcopy(element)

    rId_map = {}

    for blip in xml_copy.iter(_BLIP_TAG):
        r_embed = blip.get(f'{{{_NS_R}}}embed')
        if not r_embed:
            continue
        if r_embed in rId_map:
            blip.set(f'{{{_NS_R}}}embed', rId_map[r_embed])
            continue
        try:
            img_part = orig_part.related_parts.get(r_embed)
            if img_part is None:
                continue
            nuevo_rId = nuevo_part.relate_to(img_part, _REL_IMAGE)
        except Exception:
            try:
                from docx.opc.part import Part
                from docx.opc.packuri import PackURI
                img_part2 = orig_part.related_parts.get(r_embed)
                if img_part2 is None:
                    continue
                new_p = Part(PackURI(img_part2.partname), img_part2.content_type, img_part2._blob, nuevo_part.package)
                nuevo_rId = nuevo_part.relate_to(new_p, _REL_IMAGE)
            except Exception:
                continue
        rId_map[r_embed] = nuevo_rId
        blip.set(f'{{{_NS_R}}}embed', nuevo_rId)

    # NO eliminamos xml_copy.pPr para mantener el formato EXACTO original
    _WNS_SECTPR = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr'
    sect_pr = doc_nuevo.element.body.find(_WNS_SECTPR)
    if sect_pr is not None:
        sect_pr.addprevious(xml_copy)
    else:
        doc_nuevo.element.body.append(xml_copy)
        
    return xml_copy

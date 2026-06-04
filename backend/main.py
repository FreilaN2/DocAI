import os
import re
import sys
import shutil
import json
import uuid
import subprocess
from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks, Query, Depends
from fastapi.responses import FileResponse, StreamingResponse
import time
from datetime import datetime
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from pydantic import BaseModel
from typing import Optional, AsyncGenerator
import io
from google.oauth2 import id_token
from google.auth.transport import requests as google_requests

# Importar motores de procesamiento y base de datos
from core.database import init_db, get_db
from core.apa_rules import procesar_con_reglas, clasificar_parrafo_reglas
from core.apa_ai import procesar_con_ia, procesar_con_ia_stream
from core.auth import get_password_hash, verify_password, create_access_token
from core.models import User, Plan, TokenBalance
from core.token_service import (
    get_available_tokens,
    consume_tokens,
    groq_tokens_to_docai,
    assign_monthly_tokens,
    add_extra_tokens,
)
from core.paypal import create_order, capture_order
from sqlalchemy.orm import Session
from sqlalchemy import text
from fastapi.security import OAuth2PasswordBearer
from jose import JWTError, jwt
from core.auth import SECRET_KEY, ALGORITHM
import logging
from functools import lru_cache

# ═══════════════════════════════════════════════════════════
# CONFIGURACIÓN GLOBAL (definida al inicio, antes de usarse)
# ═══════════════════════════════════════════════════════════

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# FIX #1: BASE_DIR y directorios definidos ANTES de startup_event para evitar
# el NameError que ocurría porque startup_event los usaba antes de que se definieran.
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
PROCESSED_DIR = os.path.join(BASE_DIR, "processed")
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(PROCESSED_DIR, exist_ok=True)

# ═══════════════════════════════════════════════════════════
# CONSTANTES APA PRE-CALCULADAS (evita recalcular en cada request)
# ═══════════════════════════════════════════════════════════

LETTER_PAGE = {"width": Inches(8.5), "height": Inches(11.0)}
FUENTES_APA = {
    "Times New Roman": {"tamano": 12, "familia": "serif"},
    "Georgia": {"tamano": 11, "familia": "serif"},
    "Computer Modern": {"tamano": 10, "familia": "serif"},
    "Calibri": {"tamano": 11, "familia": "sans-serif"},
    "Arial": {"tamano": 11, "familia": "sans-serif"},
    "Lucida Sans Unicode": {"tamano": 10, "familia": "sans-serif"},
}
DEFAULT_APA_FONT = "Times New Roman"

# FIX #9: Mapa pre-normalizado (lowercase → nombre original) para validar fuentes
# en O(1) sin iterar ni hacer lower() por cada entrada en cada request.
_FUENTES_APA_LOWER: dict[str, str] = {k.lower(): k for k in FUENTES_APA}

NORMAS_APA = {
    "6ta": {
        "fuente": DEFAULT_APA_FONT,
        "tamano": FUENTES_APA[DEFAULT_APA_FONT]["tamano"],
        "interlineado": 2.0,
        "margen": 1.0,
        "page_width": LETTER_PAGE["width"],
        "page_height": LETTER_PAGE["height"],
        "sangria_primera_linea": 0.5,
        "sangria_francesa": 0.5,
        "títulos": {
            "N1": {"bold": True, "italic": False, "align": "center"},
            "N2": {"bold": True, "italic": False, "align": "left"},
            "N3": {"bold": True, "italic": False, "align": "left"},
            "N4": {"bold": True, "italic": False, "align": "indent"},
            "N5": {"bold": True, "italic": True, "align": "indent"}
        }
    },
    "7ma": {
        "fuente": DEFAULT_APA_FONT,
        "tamano": FUENTES_APA[DEFAULT_APA_FONT]["tamano"],
        "interlineado": 2.0,
        "margen": 1.0,
        "page_width": LETTER_PAGE["width"],
        "page_height": LETTER_PAGE["height"],
        "sangria_primera_linea": 0.5,
        "sangria_francesa": 0.5,
        "títulos": {
            "N1": {"bold": True, "italic": False, "align": "center"},
            "N2": {"bold": True, "italic": False, "align": "left"},
            "N3": {"bold": True, "italic": True, "align": "left"},
            "N4": {"bold": True, "italic": False, "align": "indent"},
            "N5": {"bold": True, "italic": True, "align": "indent"}
        }
    }
}

# FIX #8: Conjunto de palabras no-mayúsculas como constante de módulo (no se recrea en cada llamada)
_PALABRAS_NO_MAYUSCULAS = frozenset({
    "a", "ante", "bajo", "con", "contra", "de", "del", "desde",
    "durante", "e", "el", "la", "las", "los", "para", "por",
    "sin", "sobre", "y", "o", "u", "en", "al", "aun"
})

# FIX #6 y #15: Expresiones regulares compiladas una sola vez al cargar el módulo.
# Compilar dentro de las funciones crea un nuevo objeto regex en cada llamada.
_RE_NORMALIZAR_CAT_CHARS = re.compile(r"[^A-Z0-9_]+")
_RE_NORMALIZAR_CAT_SPACES = re.compile(r"\s+")
_RE_NORMALIZAR_TEXTO_CHARS = re.compile(r"[^a-z0-9 ]+")
_RE_AUTOR_MATCH = re.compile(r'^\s*([A-ZÁÉÍÓÚÑÜÇ][\wÁÉÍÓÚÑÜÇ\'-]+)', re.IGNORECASE)
_RE_SAFE_FILENAME = re.compile(r"[^A-Za-z0-9.-]")
_RE_SAFE_BASENAME = re.compile(r"[^A-Za-z0-9 _-]")

# Tabla de transliteración para acentos (creada una sola vez)
_TRANS_ACENTOS = str.maketrans("áéíóúüñç", "aeiouunc")

# ═══════════════════════════════════════════════════════════
# ALMACENAMIENTO EN MEMORIA CON TAMAÑO LIMITADO
# ═══════════════════════════════════════════════════════════

# FIX #2 y #3: Los dicts de storage original eran dicts sin límite ni expiración
# real (memory leak). Se reemplazan por clases con TTL y límite de entradas.
# El LRU de Python no permite TTL, así que usamos un dict simple con timestamp.

_MAX_STORAGE_ENTRIES = 500  # Límite máximo de entradas simultáneas

class _TTLStorage:
    """Dict con TTL y límite de tamaño para evitar memory leaks."""

    def __init__(self, ttl_seconds: int = 86400, max_size: int = _MAX_STORAGE_ENTRIES):
        self._data: dict[str, tuple] = {}  # key → (valor, timestamp)
        self._ttl = ttl_seconds
        self._max = max_size

    def set(self, key: str, value) -> None:
        self._evict()
        if len(self._data) >= self._max:
            # Eliminar la entrada más antigua si se alcanza el límite
            oldest = min(self._data, key=lambda k: self._data[k][1])
            del self._data[oldest]
        self._data[key] = (value, time.time())

    def get(self, key: str, default=None):
        entry = self._data.get(key)
        if entry is None:
            return default
        value, ts = entry
        if time.time() - ts > self._ttl:
            del self._data[key]
            return default
        return value

    def pop(self, key: str, default=None):
        entry = self._data.pop(key, None)
        if entry is None:
            return default
        return entry[0]

    def __contains__(self, key: str) -> bool:
        return self.get(key) is not None

    def _evict(self) -> None:
        """Elimina entradas expiradas."""
        now = time.time()
        expired = [k for k, (_, ts) in self._data.items() if now - ts > self._ttl]
        for k in expired:
            del self._data[k]


storage = _TTLStorage(ttl_seconds=86400)        # file_id → ruta del archivo
upload_storage = _TTLStorage(ttl_seconds=3600)  # upload_id → (ruta, nombre)

# ═══════════════════════════════════════════════════════════
# CACHÉ DE RUTA DEL FRONTEND (evita os.path.exists() en cada request)
# ═══════════════════════════════════════════════════════════

# FIX #12: El endpoint catchall ejecutaba dos os.path.exists() por cada request
# para determinar qué directorio usar. Con lru_cache se calcula una sola vez.
@lru_cache(maxsize=1)
def _get_frontend_dir() -> str:
    cpanel = "/home2/teleredt/public_html/docai.teleredtv.com"
    local  = os.path.join(BASE_DIR, "dist")
    return cpanel if os.path.exists(os.path.join(cpanel, "index.html")) else local


# ═══════════════════════════════════════════════════════════
# APLICACIÓN FASTAPI
# ═══════════════════════════════════════════════════════════

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://127.0.0.1:5173",
        "https://docai.teleredtv.com",
        "http://docai.teleredtv.com",
        "https://*.up.railway.app",
        "https://docai-production-6334.up.railway.app",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

oauth2_scheme          = OAuth2PasswordBearer(tokenUrl="login")
oauth2_scheme_optional = OAuth2PasswordBearer(tokenUrl="login", auto_error=False)

# ═══════════════════════════════════════════════════════════
# AUTENTICACIÓN — LÓGICA CENTRALIZADA (evita duplicación)
# ═══════════════════════════════════════════════════════════

# FIX #11: get_current_user y get_optional_current_user tenían lógica JWT duplicada.
# Se extrae a _decode_user_from_token() y ambas la reutilizan.

def _decode_user_from_token(token: str, db: Session) -> Optional[User]:
    """Decodifica el JWT y retorna el User, o None si es inválido."""
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email: str = payload.get("sub")
        if email:
            return db.query(User).filter(User.email == email).first()
    except JWTError:
        pass
    return None


def get_current_user(
    token: str = Depends(oauth2_scheme),
    db: Session = Depends(get_db),
) -> User:
    user = _decode_user_from_token(token, db)
    if user is None:
        raise HTTPException(
            status_code=401,
            detail="No se pudieron validar las credenciales",
            headers={"WWW-Authenticate": "Bearer"},
        )
    return user


def get_optional_current_user(
    token: str = Depends(oauth2_scheme_optional),
    db: Session = Depends(get_db),
) -> Optional[User]:
    if not token:
        return None
    return _decode_user_from_token(token, db)


# ═══════════════════════════════════════════════════════════
# EVENTOS DE INICIO
# ═══════════════════════════════════════════════════════════

@app.on_event("startup")
async def startup_event():
    logger.info("🛠️ Intentando inicializar base de datos...")
    try:
        init_db()
        logger.info("✅ Base de datos inicializada correctamente.")
    except Exception as e:
        logger.error(f"❌ Error en startup DB: {e}")

    logger.info("🖋️ Verificando fuentes personalizadas...")
    try:
        fonts_dir = os.path.join(BASE_DIR, "fonts")
        user_fonts_dir = os.path.expanduser("~/.local/share/fonts")

        if os.path.exists(fonts_dir):
            os.makedirs(user_fonts_dir, exist_ok=True)
            fuentes_instaladas = False

            for font_file in os.listdir(fonts_dir):
                if font_file.lower().endswith(('.ttf', '.otf')):
                    src = os.path.join(fonts_dir, font_file)
                    dst = os.path.join(user_fonts_dir, font_file)
                    if not os.path.exists(dst):
                        shutil.copy(src, dst)
                        logger.info(f"📥 Fuente copiada: {font_file}")
                        fuentes_instaladas = True

            if fuentes_instaladas:
                subprocess.run(
                    ["fc-cache", "-f", "-v"],
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                )
                logger.info("✅ Caché de fuentes actualizada.")
            else:
                logger.info("✅ Las fuentes ya estaban instaladas.")
        else:
            logger.warning("⚠️ No se encontró la carpeta 'fonts'.")
    except Exception as e:
        logger.error(f"❌ Error al cargar fuentes: {e}")


# ═══════════════════════════════════════════════════════════
# RUTAS DE DIAGNÓSTICO
# ═══════════════════════════════════════════════════════════

@app.get("/prueba-rapida")
def prueba_rapida():
    return {"status": "¡FastAPI está vivo y responde en milisegundos!"}


@app.get("/diagnostico-db")
def diagnostico_db(db: Session = Depends(get_db)):
    import traceback
    config_info = {
        "MYSQL_URL_presente": bool(os.getenv("MYSQL_URL")),
        "DB_HOST": os.getenv("DB_HOST", "NO_DEFINIDO"),
        "DB_PORT": os.getenv("DB_PORT", "NO_DEFINIDO"),
        "DB_USER": os.getenv("DB_USER", "NO_DEFINIDO"),
        "DB_NAME": os.getenv("DB_NAME", "NO_DEFINIDO"),
        "DB_PASS_presente": bool(os.getenv("DB_PASS")),
    }
    try:
        db.execute(text("SELECT 1"))
        init_db()
        return {
            "status": "success",
            "mensaje": "✅ Conexión a MySQL exitosa y tablas creadas/verificadas correctamente.",
            "config": config_info,
        }
    except Exception as e:
        return {
            "status": "error",
            "mensaje": "🚨 Falla al conectar con MySQL o al crear tablas",
            "error_real": str(e),
            "error_tipo": type(e).__name__,
            "traceback": traceback.format_exc(),
            "config": config_info,
        }


@app.get("/health")
def health_check():
    return {"status": "ok"}


# ═══════════════════════════════════════════════════════════
# LIMPIEZA DE ARCHIVOS ANTIGUOS
# ═══════════════════════════════════════════════════════════

# FIX #4: limpiar_archivos_antiguos era síncrona y bloqueaba el event loop
# al hacer I/O de disco. Se convierte en async y usa asyncio para no bloquear.
async def limpiar_archivos_antiguos():
    """Elimina archivos con más de 24 horas de antigüedad (no bloquea el event loop)."""
    import asyncio

    ahora = time.time()
    umbral = 86400  # 24 horas en segundos

    def _do_cleanup():
        for carpeta in [UPLOAD_DIR, PROCESSED_DIR]:
            if not os.path.exists(carpeta):
                continue
            for archivo in os.listdir(carpeta):
                ruta = os.path.join(carpeta, archivo)
                if os.path.isfile(ruta) and (ahora - os.path.getmtime(ruta)) > umbral:
                    try:
                        os.remove(ruta)
                        logger.info(f"🧹 Limpieza: {archivo} eliminado")
                    except Exception as e:
                        logger.error(f"❌ Error limpiando {archivo}: {e}")

    # Ejecutar el I/O de disco en un thread para no bloquear el event loop
    await asyncio.get_event_loop().run_in_executor(None, _do_cleanup)


# ═══════════════════════════════════════════════════════════
# HELPERS APA — OPTIMIZADOS
# ═══════════════════════════════════════════════════════════

def validar_fuente_apa(nombre_fuente: str) -> str:
    # FIX #9: lookup O(1) con dict pre-normalizado en lugar de iterar y hacer lower()
    if not nombre_fuente:
        return DEFAULT_APA_FONT
    return _FUENTES_APA_LOWER.get(nombre_fuente.strip().lower(), DEFAULT_APA_FONT)


def formatear_titulo_apa(texto: str) -> str:
    # FIX #8: _PALABRAS_NO_MAYUSCULAS es frozenset de módulo, no se recrea
    if not texto:
        return texto.strip()
    partes = texto.strip().split()
    resultado = []
    for idx, palabra in enumerate(partes):
        lower = palabra.strip().lower()
        if idx == 0 or lower not in _PALABRAS_NO_MAYUSCULAS:
            resultado.append(lower.capitalize())
        else:
            resultado.append(lower)
    return " ".join(resultado)


def formatear_titulo_apa_sentence_case(texto: str) -> str:
    if not texto:
        return texto.strip()
    t = texto.strip().lower()
    return (t[0].upper() + t[1:]) if len(t) > 1 else t.upper()


def formatear_titulo_por_nivel(texto: str, nivel: str, edicion: str) -> str:
    if edicion == "6ta" and nivel in {"N3", "N4", "N5"}:
        return formatear_titulo_apa_sentence_case(texto)
    return formatear_titulo_apa(texto)


def _normalizar_categoria(categoria: str) -> str:
    # FIX #6: usa regex compilados (_RE_*) en lugar de compilar en cada llamada
    if not categoria:
        return "PARRAFO_NORMAL"

    cat = _RE_NORMALIZAR_CAT_CHARS.sub("", categoria.strip().upper().replace(" ", "_"))

    if cat.startswith("TITULO") and "_N" not in cat:
        cat = "TITULO_N1"

    if cat.startswith("REFERENCIA"):
        return "REFERENCIA"
    if cat in {"CITA_LARGA", "BLOQUE_CITA"}:
        return cat
    if cat in {"TITULO_N1", "TITULO_N2", "TITULO_N3", "TITULO_N4", "TITULO_N5", "PARRAFO_NORMAL"}:
        return cat

    return "PARRAFO_NORMAL"


def _normalizar_texto_para_encabezado(texto: str) -> str:
    # FIX #15: usa tabla de transliteración y regex compilado de módulo
    t = texto.strip().lower().translate(_TRANS_ACENTOS)
    return _RE_NORMALIZAR_TEXTO_CHARS.sub("", t).strip()


# Conjunto pre-calculado de encabezados de referencias válidos
_ENCABEZADOS_REFERENCIAS = frozenset({
    "referencias",
    "referencia",
    "referencias bibliograficas",
    "referencia bibliografica",
    "references bibliograficas",
    "bibliografia",
    "bibliografias",
    "bibliograficas",
    "bibliografica",
})

_CONTINUACIONES_REFERENCIAS = frozenset({
    "bibliografia",
    "bibliografias",
    "bibliograficas",
    "bibliografica",
    "bibliografico",
})


def _es_encabezado_referencias(texto: str) -> bool:
    base = _normalizar_texto_para_encabezado(texto)
    if base in _ENCABEZADOS_REFERENCIAS:
        return True
    return ("referencia" in base and "bibliograf" in base) or base.startswith("referencia")


def _es_continuacion_encabezado_referencias(texto: str) -> bool:
    base = _normalizar_texto_para_encabezado(texto)
    return base in _CONTINUACIONES_REFERENCIAS or base.startswith("bibliograf")


def _ordenar_referencia_por_autor(texto: str) -> str:
    # FIX #7: usa regex compilado de módulo
    m = _RE_AUTOR_MATCH.match(texto.strip())
    if m:
        return m.group(1).lower()
    return _RE_NORMALIZAR_TEXTO_CHARS.sub("", texto.strip().lower())


# ═══════════════════════════════════════════════════════════
# ESTILOS DE PÁRRAFO APA
# ═══════════════════════════════════════════════════════════

def configurar_parrafo_estilo(paragraph, categoria, reglas, body_text: str = None):
    categoria = _normalizar_categoria(categoria)
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(0)
    pf.keep_together     = True
    pf.first_line_indent = Inches(0)
    pf.left_indent       = Inches(0)
    paragraph.alignment  = WD_ALIGN_PARAGRAPH.LEFT

    if categoria.startswith("TITULO"):
        nivel   = categoria.split("_")[-1] if "_" in categoria else "N1"
        edicion = reglas.get("edicion", "7ma")
        paragraph.text   = formatear_titulo_por_nivel(paragraph.text, nivel, edicion)
        pf.space_before  = Pt(12)
        pf.space_after   = Pt(0)
        pf.first_line_indent = Inches(0)
        pf.left_indent       = Inches(0)

        heading_map = {"N1": "Heading 1", "N2": "Heading 2"}
        if edicion == "6ta" and nivel in {"N3", "N4", "N5"}:
            paragraph.style  = "Normal"
            pf.left_indent   = Inches(reglas["sangria_primera_linea"])
            pf.first_line_indent = Inches(0)
        else:
            paragraph.style = heading_map.get(nivel, "Heading 1")
            if nivel in {"N3", "N4", "N5"}:
                pf.left_indent = Inches(reglas["sangria_primera_linea"])

        bold_italic_align = {
            "N1": (True,  False, WD_ALIGN_PARAGRAPH.CENTER),
            "N2": (True,  False, WD_ALIGN_PARAGRAPH.LEFT),
            "N3": (True,  False, WD_ALIGN_PARAGRAPH.LEFT),
            "N4": (True,  True,  WD_ALIGN_PARAGRAPH.LEFT),
            "N5": (False, True,  WD_ALIGN_PARAGRAPH.LEFT),
        }
        bold, italic, align = bold_italic_align.get(nivel, (True, False, WD_ALIGN_PARAGRAPH.LEFT))
        paragraph.alignment = align

        if body_text and edicion == "6ta" and nivel in {"N3", "N4", "N5"}:
            heading_text = formatear_titulo_por_nivel(paragraph.text, nivel, edicion)
            if not heading_text.strip().endswith('.'):
                heading_text = heading_text.rstrip() + '.'
            paragraph.text = ""
            hr = paragraph.add_run(heading_text)
            hr.bold, hr.italic, hr.font.name, hr.font.size = bold, italic, reglas["fuente"], Pt(reglas["tamano"])
            paragraph.add_run(" ")
            br = paragraph.add_run(body_text.strip())
            br.bold, br.italic, br.font.name, br.font.size = False, False, reglas["fuente"], Pt(reglas["tamano"])
            return

        if edicion == "6ta" and nivel in {"N3", "N4", "N5"} and not paragraph.text.strip().endswith('.'):
            paragraph.text = paragraph.text.rstrip() + '.'

        for run in paragraph.runs:
            run.bold, run.italic = bold, italic
            run.font.name, run.font.size = reglas["fuente"], Pt(reglas["tamano"])
        return

    if categoria in {"CITA_LARGA", "BLOQUE_CITA"}:
        paragraph.text   = paragraph.text.strip().strip('"""„»')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.left_indent   = Inches(reglas["sangria_primera_linea"])
        pf.first_line_indent = Inches(0)
        for run in paragraph.runs:
            run.font.name, run.font.size = reglas["fuente"], Pt(reglas["tamano"])
        return

    if categoria == "REFERENCIA":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Inches(-reglas["sangria_francesa"])
        pf.left_indent       = Inches(reglas["sangria_francesa"])
        pf.keep_together     = True
        pf.space_before      = Pt(0)
        pf.space_after       = Pt(0)
        for run in paragraph.runs:
            run.font.name, run.font.size = reglas["fuente"], Pt(reglas["tamano"])
        return

    # PARRAFO_NORMAL (default)
    pf.first_line_indent = Inches(reglas["sangria_primera_linea"])
    paragraph.alignment  = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        run.font.name, run.font.size = reglas["fuente"], Pt(reglas["tamano"])


# ═══════════════════════════════════════════════════════════
# UTILIDADES DE DOCUMENTO WORD
# ═══════════════════════════════════════════════════════════

def añadir_marca_de_agua(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Generado con DocAI Free — Formateador Automático APA")
        run.font.size = Pt(10)
        run.font.name = "Arial"


def _insertar_tabla_de_contenidos(doc):
    paragraph = doc.add_paragraph()
    fld_simple = OxmlElement('w:fldSimple')
    fld_simple.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    fld_simple.set(qn('w:dirty'), 'true')
    paragraph._p.append(fld_simple)


def _configurar_encabezado_paginas(doc):
    for section in doc.sections:
        section.header.is_linked_to_previous = False
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run()
        fld_simple = OxmlElement('w:fldSimple')
        fld_simple.set(qn('w:instr'), 'PAGE \\* MERGEFORMAT')
        run._r.append(fld_simple)


def _force_update_fields(doc):
    settings = getattr(doc, 'settings', None)
    if not settings:
        return
    element = getattr(settings, 'element', None)
    if not element:
        return
    update = OxmlElement('w:updateFields')
    update.set(qn('w:val'), 'true')
    element.append(update)


def _get_soffice_path() -> Optional[str]:
    """Detecta la ruta de LibreOffice. Resultado cacheado implícitamente por lru_cache abajo."""
    path_in_env = shutil.which("soffice")
    if path_in_env:
        return path_in_env
    for p in [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]:
        if os.path.exists(p):
            return p
    return None

# FIX extra: cachear la ruta de soffice para no llamar shutil.which() en cada conversión PDF
_get_soffice_path_cached = lru_cache(maxsize=1)(_get_soffice_path)


# ═══════════════════════════════════════════════════════════
# MODELOS PYDANTIC
# ═══════════════════════════════════════════════════════════

class ParrafoCorregido(BaseModel):
    texto: str
    categoria: str

class DatosFinales(BaseModel):
    edicion: str
    parrafos: list[ParrafoCorregido]  # FIX #14: typing moderno (Python 3.9+)
    filename: str
    plan: str = "free"
    incluir_indice: bool = False
    formato: str = "docx"
    fuente: str = DEFAULT_APA_FONT

class UserCreate(BaseModel):
    firstName: str
    lastName: str
    email: str
    phone: str
    country: str
    password: str

class UserLogin(BaseModel):
    email: str
    password: str

class GoogleAuthRequest(BaseModel):
    token: str

class SuscripcionRequest(BaseModel):
    months: int

class ConfirmarPagoRequest(BaseModel):
    order_id: str
    months: int

class PackRequest(BaseModel):
    pack_id: int

class ConfirmarPackRequest(BaseModel):
    order_id: str
    pack_id: int

class VerifyBinanceRequest(BaseModel):
    order_id: str
    type: str
    item_id: int


# ═══════════════════════════════════════════════════════════
# PRECIOS (constantes de módulo, no inline en endpoints)
# ═══════════════════════════════════════════════════════════

SUBSCRIPTION_PRICES = {1: 5.00, 3: 14.00, 6: 25.00, 12: 45.00}
TOKENS_PER_MONTH_PRO = 500


# ═══════════════════════════════════════════════════════════
# ENDPOINTS DE AUTENTICACIÓN
# ═══════════════════════════════════════════════════════════

@app.post("/register")
def register(user_data: UserCreate, db: Session = Depends(get_db)):
    if db.query(User).filter(User.email == user_data.email).first():
        raise HTTPException(status_code=400, detail="El correo electrónico ya está registrado.")
    if user_data.phone and db.query(User).filter(User.phone == user_data.phone).first():
        raise HTTPException(status_code=400, detail="Este número de teléfono ya está asociado a otra cuenta.")

    new_user = User(
        first_name=user_data.firstName,
        last_name=user_data.lastName,
        email=user_data.email,
        phone=user_data.phone,
        country=user_data.country,
        password_hash=get_password_hash(user_data.password),
        plan_id=1,
    )
    db.add(new_user)
    db.commit()
    db.refresh(new_user)

    access_token = create_access_token(data={"sub": new_user.email})
    return {
        "status": "success",
        "access_token": access_token,
        "token_type": "bearer",
        "user": {
            "id": new_user.id,
            "email": new_user.email,
            "firstName": new_user.first_name,
            "lastName": new_user.last_name,
            "phone": new_user.phone,
            "country": new_user.country,
            "plan": new_user.plan.name if new_user.plan else "free",
            "createdAt": new_user.created_at.isoformat() if getattr(new_user, 'created_at', None) else None,
            "lastLoginAt": None,
        },
    }


@app.post("/login")
def login(credentials: UserLogin, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == credentials.email).first()
    if not user or not verify_password(credentials.password, user.password_hash):
        raise HTTPException(status_code=401, detail="Correo o contraseña incorrectos.")

    user.last_login_at = datetime.utcnow()
    db.commit()

    access_token = create_access_token(data={"sub": user.email})
    return {
        "status": "success",
        "access_token": access_token,
        "token_type": "bearer",
        "user": {
            "id": user.id,
            "email": user.email,
            "firstName": user.first_name,
            "lastName": user.last_name,
            "phone": user.phone,
            "country": user.country,
            "plan": user.plan.name if user.plan else "free",
            "createdAt": user.created_at.isoformat() if getattr(user, 'created_at', None) else None,
            "lastLoginAt": user.last_login_at.isoformat() if getattr(user, 'last_login_at', None) else None,
        },
    }


@app.post("/auth/google")
def auth_google(data: GoogleAuthRequest, db: Session = Depends(get_db)):
    try:
        client_id = os.getenv("GOOGLE_CLIENT_ID")
        id_info = id_token.verify_oauth2_token(data.token, google_requests.Request(), client_id)
        email      = id_info.get("email")
        first_name = id_info.get("given_name", "Google")
        last_name  = id_info.get("family_name", "User")

        user = db.query(User).filter(User.email == email).first()
        if not user:
            user = User(
                first_name=first_name,
                last_name=last_name,
                email=email,
                phone=None,
                country="US",
                password_hash=get_password_hash(os.urandom(24).hex()),
                plan_id=1,
            )
            db.add(user)
            db.commit()
            db.refresh(user)

        user.last_login_at = datetime.utcnow()
        db.commit()

        access_token = create_access_token(data={"sub": user.email})
        return {
            "status": "success",
            "access_token": access_token,
            "token_type": "bearer",
            "user": {
                "id": user.id,
                "email": user.email,
                "firstName": user.first_name,
                "lastName": user.last_name,
                "phone": user.phone,
                "country": user.country,
                "plan": user.plan.name if user.plan else "free",
            },
        }
    except ValueError:
        raise HTTPException(status_code=400, detail="Token de Google inválido o expirado.")
    except Exception as e:
        logger.error(f"Error en Google Auth: {e}")
        raise HTTPException(status_code=500, detail=f"Error interno: {e}")


# ═══════════════════════════════════════════════════════════
# ENDPOINTS DE PROCESAMIENTO APA
# ═══════════════════════════════════════════════════════════

@app.post("/upload-documento/")
async def upload_documento(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    current_user: User = Depends(get_optional_current_user),
):
    background_tasks.add_task(limpiar_archivos_antiguos)

    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Solo se aceptan archivos .docx")

    contents   = await file.read()
    upload_id  = str(uuid.uuid4())
    safe_name  = _RE_SAFE_FILENAME.sub("_", file.filename) if file.filename else "upload.docx"
    input_path = os.path.join(UPLOAD_DIR, f"{upload_id}_{safe_name}")

    with open(input_path, "wb") as f:
        f.write(contents)

    upload_storage.set(upload_id, (input_path, safe_name))
    logger.info(f"📤 Upload #{upload_id}: {safe_name} ({len(contents)} bytes)")
    return {"upload_id": upload_id, "filename": safe_name}


@app.get("/procesar-apa/stream")
async def procesar_apa_stream(
    upload_id: str = Query(...),
    edicion: str   = Query("7ma"),
    plan: str      = Query("free"),
    token: str     = Query(None),
    db: Session    = Depends(get_db),
):
    # FIX #5: Validación JWT reutiliza _decode_user_from_token (sin duplicar lógica)
    current_user = None
    if token and token != "null":
        current_user = _decode_user_from_token(token, db)

    entry = upload_storage.get(upload_id)
    if entry is None:
        raise HTTPException(status_code=404, detail="upload_id no encontrado. Sube el archivo primero.")
    input_path, filename = entry

    if plan == "pro":
        if not current_user:
            raise HTTPException(status_code=401, detail="Debes iniciar sesión para usar DocAI Pro.")
        if get_available_tokens(current_user.id, db)["total"] <= 0:
            raise HTTPException(status_code=402, detail="Sin tokens disponibles.")

    try:
        with open(input_path, "rb") as f:
            doc = Document(io.BytesIO(f.read()))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"No se pudo leer el .docx: {e}")

    async def event_generator() -> AsyncGenerator[str, None]:
        try:
            if plan == "pro":
                async for evento in procesar_con_ia_stream(doc.paragraphs):
                    if evento.get("tipo") == "finalizado":
                        consume_tokens(current_user.id, evento.get("groq_tokens", 0), filename, db)
                        try:
                            os.remove(input_path)
                            upload_storage.pop(upload_id)
                        except Exception:
                            pass
                    yield f"data: {json.dumps(evento, ensure_ascii=False)}\n\n"
            else:
                resultado = procesar_con_reglas(doc.paragraphs)
                yield f"data: {json.dumps({'tipo': 'inicio', 'total_lotes': 1, 'progreso': 0, 'modelo': 'reglas'}, ensure_ascii=False)}\n\n"
                yield f"data: {json.dumps({'tipo': 'finalizado', 'progreso': 100, 'stats': resultado['stats'], 'detalles': resultado['detalles'], 'groq_tokens': 0}, ensure_ascii=False)}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'tipo': 'error', 'mensaje': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.post("/procesar-apa/")
async def procesar_documento(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    edicion: str = Form("7ma"),
    plan: str = Form("free"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_optional_current_user),
):
    background_tasks.add_task(limpiar_archivos_antiguos)

    input_path = os.path.join(UPLOAD_DIR, file.filename)
    contents   = await file.read()
    with open(input_path, "wb") as f:
        f.write(contents)

    if not os.path.exists(input_path):
        logger.error(f"❌ El archivo no se creó en {input_path}")
        raise HTTPException(status_code=500, detail="Error al guardar el archivo")

    logger.info(f"📂 {input_path} ({os.path.getsize(input_path)} bytes)")

    try:
        doc = Document(io.BytesIO(contents))
    except Exception as e:
        logger.error(f"❌ Error al abrir .docx: {e}")
        raise HTTPException(status_code=400, detail=f"No se pudo procesar el .docx: {e}")

    logger.info(f"🚀 Procesando: {file.filename} (Plan: {plan}, Edición: {edicion})")

    if plan == "pro":
        balance = get_available_tokens(current_user.id, db)
        if balance["total"] <= 0:
            raise HTTPException(status_code=402, detail="No tienes tokens disponibles.")

        logger.info("🤖 Usando motor IA (Groq Llama 3.3)")
        resultado    = procesar_con_ia(doc.paragraphs)
        groq_tokens  = resultado.get('groq_tokens', 0)
        consume_tokens(current_user.id, groq_tokens, file.filename, db)
        logger.info(f"💡 Tokens consumidos: {groq_tokens_to_docai(groq_tokens)} DocAI ({groq_tokens} Groq)")
    else:
        logger.info("⚖️ Usando motor de reglas (Free)")
        resultado = procesar_con_reglas(doc.paragraphs)
        resultado["groq_tokens"] = 0

    logger.info(f"✅ Párrafos procesados: {len(resultado['detalles'])}")
    return {
        "status": "success",
        "plan": plan,
        "resumen": resultado["stats"],
        "detalles": resultado["detalles"],
        "tokens_consumed": groq_tokens_to_docai(resultado.get("groq_tokens", 0)),
    }


@app.get("/tokens/balance")
async def mis_tokens(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    return {"status": "success", **get_available_tokens(current_user.id, db)}


@app.post("/generar-final/")
async def generar_final(
    datos: DatosFinales,
    current_user: User = Depends(get_current_user),
):
    base_name     = os.path.splitext(datos.filename)[0]
    safe_base     = _RE_SAFE_BASENAME.sub("_", base_name).strip()
    unique_suffix = uuid.uuid4().hex
    out_name      = f"FINAL_{datos.edicion}_{safe_base}_{unique_suffix}"
    out_docx      = os.path.join(PROCESSED_DIR, out_name + ".docx")
    output_path   = out_docx

    doc   = Document()
    reglas = dict(NORMAS_APA.get(datos.edicion, NORMAS_APA["7ma"]))
    reglas["edicion"] = datos.edicion
    if datos.edicion == "6ta":
        reglas["fuente"] = DEFAULT_APA_FONT
        reglas["tamano"] = FUENTES_APA[DEFAULT_APA_FONT]["tamano"]
    else:
        reglas["fuente"] = validar_fuente_apa(datos.fuente)
        reglas["tamano"] = FUENTES_APA[reglas["fuente"]]["tamano"]

    for section in doc.sections:
        section.page_width = section.page_height = None  # reset
        section.page_width    = LETTER_PAGE["width"]
        section.page_height   = LETTER_PAGE["height"]
        m = Inches(reglas["margen"])
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = m

    ns = doc.styles["Normal"]
    ns.font.name = reglas["fuente"]
    ns.font.size = Pt(reglas["tamano"])
    nspf = ns.paragraph_format
    nspf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    nspf.space_before = nspf.space_after = Pt(0)
    nspf.first_line_indent = Inches(reglas["sangria_primera_linea"])
    nspf.left_indent = Inches(0)

    _configurar_encabezado_paginas(doc)

    # --- Índice ---
    if datos.incluir_indice and datos.plan == "pro":
        headings = []
        for p in datos.parrafos:
            cat = _normalizar_categoria(p.categoria)
            if "TITULO" in cat:
                try:
                    nivel = int(cat.split("_")[-1].replace("N", ""))
                except Exception:
                    nivel = 1
                headings.append((min(max(nivel, 1), 3), p.texto.strip()))

        title = doc.add_paragraph("Índice")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.bold, run.font.name, run.font.size = True, reglas["fuente"], Pt(16)

        if headings:
            _insertar_tabla_de_contenidos(doc)
            nota = doc.add_paragraph("Actualiza los campos en Word para obtener los números de página reales.")
            nota.italic = True
            nota.paragraph_format.space_before = Pt(4)
            nota.paragraph_format.space_after  = Pt(8)
        else:
            w = doc.add_paragraph("No se detectaron títulos válidos.")
            w.italic = True

        doc.add_page_break()

    # --- Cuerpo del documento ---
    reference_started  = False
    paragraph_counter  = 0
    reference_buffer: list[ParrafoCorregido] = []
    i = 0
    parrafos = datos.parrafos  # referencia local evita lookup de atributo en cada iteración

    while i < len(parrafos):
        p   = parrafos[i]
        cat = _normalizar_categoria(p.categoria)

        # Títulos N3-N5 en 6ta edición: fusionar con el párrafo siguiente
        if datos.edicion == "6ta" and cat in {"TITULO_N3", "TITULO_N4", "TITULO_N5"} and i + 1 < len(parrafos):
            if _normalizar_categoria(parrafos[i + 1].categoria) == "PARRAFO_NORMAL":
                paragraph = doc.add_paragraph()
                configurar_parrafo_estilo(paragraph, cat, reglas, body_text=parrafos[i + 1].texto.strip())
                paragraph_counter += 1
                i += 2
                continue

        # Detectar sección de referencias
        if not reference_started and _es_encabezado_referencias(p.texto):
            if paragraph_counter > 0:
                doc.add_page_break()

            heading_texto = p.texto.strip()
            if i + 1 < len(parrafos) and _es_continuacion_encabezado_referencias(parrafos[i + 1].texto):
                heading_texto += " " + parrafos[i + 1].texto.strip()
                i += 1

            ref_h = doc.add_paragraph(heading_texto)
            ref_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in ref_h.runs:
                run.bold = datos.edicion != "6ta"
                run.font.name, run.font.size = reglas["fuente"], Pt(reglas["tamano"])
            ref_h.paragraph_format.space_before = Pt(12)
            ref_h.paragraph_format.space_after  = Pt(12)
            reference_started = True
            paragraph_counter += 1
            i += 1
            continue

        if cat == "REFERENCIA" and not reference_started:
            if paragraph_counter > 0:
                doc.add_page_break()
            ref_h = doc.add_paragraph("Referencias")
            ref_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in ref_h.runs:
                run.bold = datos.edicion != "6ta"
                run.font.name, run.font.size = reglas["fuente"], Pt(reglas["tamano"])
            ref_h.paragraph_format.space_before = Pt(12)
            ref_h.paragraph_format.space_after  = Pt(12)
            reference_started = True

        if reference_started and cat == "REFERENCIA":
            reference_buffer.append(p)
            i += 1
            continue

        # Volcar buffer de referencias ordenado antes de continuar con párrafos normales
        if reference_started and reference_buffer:
            for ref in sorted(reference_buffer, key=lambda r: _ordenar_referencia_por_autor(r.texto)):
                ph = doc.add_paragraph(ref.texto)
                configurar_parrafo_estilo(ph, ref.categoria, reglas)
                paragraph_counter += 1
            reference_buffer = []

        ph = doc.add_paragraph(p.texto)
        configurar_parrafo_estilo(ph, cat, reglas)
        paragraph_counter += 1
        i += 1

    # Volcar referencias restantes (fin de documento)
    if reference_buffer:
        for ref in sorted(reference_buffer, key=lambda r: _ordenar_referencia_por_autor(r.texto)):
            ph = doc.add_paragraph(ref.texto)
            configurar_parrafo_estilo(ph, ref.categoria, reglas)
            paragraph_counter += 1

    if datos.incluir_indice and datos.plan == "pro":
        _force_update_fields(doc)

    if datos.plan == "free":
        añadir_marca_de_agua(doc)

    try:
        doc.save(out_docx)
    except PermissionError as e:
        logger.error(f"❌ Permiso denegado al guardar DOCX: {e}")
        raise HTTPException(status_code=500, detail="No se pudo guardar el archivo DOCX.")
    except Exception as e:
        logger.error(f"❌ Error al guardar DOCX: {e}")
        raise HTTPException(status_code=500, detail="No se pudo generar el archivo DOCX.")

    # --- Conversión a PDF (solo Pro) ---
    if datos.formato.lower() == "pdf":
        if datos.plan != "pro":
            raise HTTPException(status_code=403, detail="PDF exclusivo para usuarios Pro.")

        out_pdf = os.path.abspath(os.path.join(PROCESSED_DIR, out_name + ".pdf"))
        if os.path.exists(out_pdf):
            try:
                os.remove(out_pdf)
            except Exception:
                pass

        soffice = _get_soffice_path_cached()
        if not soffice:
            raise HTTPException(status_code=500, detail="LibreOffice no está instalado en el servidor.")

        try:
            import tempfile
            tmp_dir = tempfile.mkdtemp()
            result  = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", tmp_dir,
                 os.path.abspath(out_docx)],
                capture_output=True, text=True, timeout=60,
            )
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice falló: {result.stderr}")

            generated = os.path.join(tmp_dir, os.path.splitext(os.path.basename(out_docx))[0] + ".pdf")
            if not os.path.exists(generated):
                raise RuntimeError("LibreOffice no generó el PDF esperado.")

            shutil.move(generated, out_pdf)
            shutil.rmtree(tmp_dir, ignore_errors=True)
            output_path = out_pdf
            logger.info(f"✅ PDF generado: {out_pdf}")
        except subprocess.TimeoutExpired:
            raise HTTPException(status_code=500, detail="La conversión a PDF tardó demasiado.")
        except Exception as e:
            logger.error(f"❌ Error PDF: {e}")
            raise HTTPException(status_code=500, detail=f"No se pudo convertir a PDF: {e}")

    # FIX #3: Usar UUID en lugar de hash() para evitar colisiones
    file_id = str(uuid.uuid4())
    storage.set(file_id, output_path)
    return {"file_id": file_id}


@app.get("/descargar/{file_id}")
async def descargar_archivo(file_id: str):
    path = storage.get(file_id)
    if path and os.path.exists(path):
        return FileResponse(path=path, filename=os.path.basename(path))
    raise HTTPException(status_code=404, detail="No encontrado")


# ═══════════════════════════════════════════════════════════
# ENDPOINTS DE PAGO — PAYPAL
# ═══════════════════════════════════════════════════════════

@app.post("/pago/suscripcion")
async def crear_orden_suscripcion(
    data: SuscripcionRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if data.months not in SUBSCRIPTION_PRICES:
        raise HTTPException(status_code=400, detail="Duración no válida. Usa 1, 3, 6 o 12.")

    amount      = SUBSCRIPTION_PRICES[data.months]
    description = f"DocAI Pro — {data.months} mes(es) | {TOKENS_PER_MONTH_PRO} tokens/mes"
    custom_id   = f"sub:{current_user.id}:{data.months}"

    try:
        order = create_order(amount=amount, description=description, custom_id=custom_id)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Error con PayPal: {e}")

    return {
        "status": "success",
        "order_id": order["order_id"],
        "approval_url": order["approval_url"],
        "amount": amount,
        "months": data.months,
    }


@app.post("/pago/confirmar-suscripcion")
async def confirmar_suscripcion(
    data: ConfirmarPagoRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    from datetime import datetime
    from dateutil.relativedelta import relativedelta
    from core.models import Subscription

    try:
        result = capture_order(data.order_id)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Error capturando pago: {e}")

    if result.get("status") != "COMPLETED":
        raise HTTPException(status_code=402, detail="El pago no fue completado por PayPal.")

    pro_plan = db.query(Plan).filter(Plan.name == "pro").first()
    current_user.plan_id = pro_plan.id

    now = datetime.utcnow()
    sub = Subscription(
        user_id=current_user.id,
        paypal_order_id=data.order_id,
        months_paid=data.months,
        tokens_per_month=TOKENS_PER_MONTH_PRO,
        started_at=now,
        ends_at=now + relativedelta(months=data.months),
        status="active",
    )
    db.add(sub)
    db.commit()
    assign_monthly_tokens(current_user.id, TOKENS_PER_MONTH_PRO, db)

    logger.info(f"🎉 Pro activado: user={current_user.id}, meses={data.months}")
    return {
        "status": "success",
        "message": f"Suscripción Pro activada por {data.months} mes(es).",
        "tokens_assigned": TOKENS_PER_MONTH_PRO,
    }


@app.post("/pago/verify-binance")
async def verify_binance(
    data: VerifyBinanceRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    from datetime import datetime
    from dateutil.relativedelta import relativedelta
    from core.models import Subscription, BinanceTransaction, TokenPack
    from core.binance_pay import verify_binance_payment

    if data.type == 'subscription':
        if data.item_id not in SUBSCRIPTION_PRICES:
            raise HTTPException(status_code=400, detail="Duración no válida.")
        expected_amount = float(SUBSCRIPTION_PRICES[data.item_id])
    elif data.type == 'pack':
        pack = db.query(TokenPack).filter(TokenPack.id == data.item_id).first()
        if not pack:
            raise HTTPException(status_code=404, detail="Paquete no encontrado.")
        expected_amount = float(pack.price)
    else:
        raise HTTPException(status_code=400, detail="Tipo de pago no válido.")

    if db.query(BinanceTransaction).filter(BinanceTransaction.order_id == data.order_id).first():
        raise HTTPException(status_code=400, detail="Este comprobante ya fue procesado.")

    is_valid, msg = verify_binance_payment(data.order_id, expected_amount)
    if not is_valid:
        raise HTTPException(status_code=400, detail=msg)

    db.add(BinanceTransaction(
        user_id=current_user.id,
        order_id=data.order_id,
        amount=expected_amount,
        currency="USDT",
    ))

    if data.type == 'subscription':
        pro_plan = db.query(Plan).filter(Plan.name == "pro").first()
        current_user.plan_id = pro_plan.id
        now = datetime.utcnow()
        db.add(Subscription(
            user_id=current_user.id,
            paypal_order_id=f"binance_{data.order_id}",
            months_paid=data.item_id,
            tokens_per_month=TOKENS_PER_MONTH_PRO,
            started_at=now,
            ends_at=now + relativedelta(months=data.item_id),
            status="active",
        ))
        assign_monthly_tokens(current_user.id, TOKENS_PER_MONTH_PRO, db)
        message = f"Pago verificado. ¡Pro activado por {data.item_id} mes(es)!"
    else:
        add_extra_tokens(current_user.id, pack.tokens, db)
        message = f"Pago verificado. ¡+{pack.tokens} tokens añadidos!"

    db.commit()
    logger.info(f"🎉 Binance Pay: user={current_user.id}, type={data.type}, item={data.item_id}")
    return {"status": "success", "message": message}


@app.post("/pago/pack-tokens")
async def crear_orden_pack(
    data: PackRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    from core.models import TokenPack

    pack = db.query(TokenPack).filter(TokenPack.id == data.pack_id, TokenPack.is_active == True).first()
    if not pack:
        raise HTTPException(status_code=404, detail="Paquete no encontrado.")

    try:
        order = create_order(
            amount=float(pack.price),
            description=f"DocAI — {pack.name} ({pack.tokens} tokens extra)",
            custom_id=f"pack:{current_user.id}:{pack.id}",
        )
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Error con PayPal: {e}")

    return {
        "status": "success",
        "order_id": order["order_id"],
        "approval_url": order["approval_url"],
        "pack": {"name": pack.name, "tokens": pack.tokens, "price": float(pack.price)},
    }


@app.post("/pago/confirmar-pack")
async def confirmar_pack(
    data: ConfirmarPackRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    from core.models import TokenPack

    try:
        result = capture_order(data.order_id)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Error capturando pago: {e}")

    if result.get("status") != "COMPLETED":
        raise HTTPException(status_code=402, detail="El pago no fue completado por PayPal.")

    pack = db.query(TokenPack).filter(TokenPack.id == data.pack_id).first()
    if not pack:
        raise HTTPException(status_code=404, detail="Paquete no encontrado.")

    add_extra_tokens(current_user.id, pack.tokens, db)
    logger.info(f"📦 Pack: user={current_user.id}, tokens=+{pack.tokens}")
    return {
        "status": "success",
        "message": f"+{pack.tokens} tokens extra añadidos.",
        "pack": pack.name,
    }


@app.get("/packs")
async def listar_packs(db: Session = Depends(get_db)):
    from core.models import TokenPack
    packs = db.query(TokenPack).filter(TokenPack.is_active == True).all()
    return [{"id": p.id, "name": p.name, "price": float(p.price), "tokens": p.tokens} for p in packs]


# ═══════════════════════════════════════════════════════════
# SERVIR FRONTEND REACT (catchall)
# ═══════════════════════════════════════════════════════════

@app.get("/{catchall:path}")
def serve_react_app(catchall: str):
    # FIX #12: _get_frontend_dir() está cacheada con lru_cache → un solo os.path.exists() en toda la vida del proceso
    frontend = _get_frontend_dir()
    file_path = os.path.join(frontend, catchall)

    if catchall and os.path.isfile(file_path):
        return FileResponse(file_path)

    index_path = os.path.join(frontend, "index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path)

    return {
        "detail": "Error: index.html no encontrado.",
        "tried": frontend,
    }
"""
routers/apa.py
==============
Endpoints de procesamiento y generación de documentos APA.
"""

import asyncio
import json
import os
import io
import uuid
import subprocess
import tempfile
import shutil
import time
from typing import AsyncGenerator, Optional

from fastapi import APIRouter, BackgroundTasks, Depends, File, Form, HTTPException, Query, Request, UploadFile
from fastapi.responses import FileResponse, StreamingResponse
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from sqlalchemy.orm import Session

from core.database import get_db
from core.models import User
from core.apa_rules import procesar_con_reglas
from core.apa_ai import procesar_con_ia, procesar_con_ia_stream
from core.token_service import get_available_tokens, consume_tokens, groq_tokens_to_docai
from core.dependencies import get_current_user, get_optional_current_user, _decode_user_from_token
from core.schemas import DatosFinales, ParrafoCorregido
from core.storage import storage, upload_storage
from core.config import UPLOAD_DIR, PROCESSED_DIR, RE_SAFE_FILENAME, RE_SAFE_BASENAME
from core.limiter import limiter
from core.document_builder import (
    NORMAS_APA, FUENTES_APA, DEFAULT_APA_FONT, LETTER_PAGE,
    validar_fuente_apa, _normalizar_categoria,
    _es_encabezado_referencias, _es_continuacion_encabezado_referencias,
    _ordenar_referencia_por_autor, configurar_parrafo_estilo,
    _insertar_tabla_de_contenidos, _configurar_encabezado_paginas,
    _force_update_fields, añadir_marca_de_agua, _get_soffice_path_cached,
    copiar_portada_desde_original, _detectar_n_portada,
)
import logging

logger = logging.getLogger(__name__)
router = APIRouter()


# ─── Limpieza de archivos ─────────────────────────────────

async def limpiar_archivos_antiguos():
    """Elimina archivos con más de 24h de antigüedad sin bloquear el event loop."""
    ahora  = time.time()
    umbral = 86400

    def _do_cleanup():
        for carpeta in [UPLOAD_DIR, PROCESSED_DIR]:
            if not os.path.exists(carpeta):
                continue
            for archivo in os.listdir(carpeta):
                ruta = os.path.join(carpeta, archivo)
                if os.path.isfile(ruta) and (ahora - os.path.getmtime(ruta)) > umbral:
                    try:
                        os.remove(ruta)
                        logger.info(f"Limpieza: {archivo} eliminado")
                    except Exception as e:
                        logger.error(f"Error limpiando {archivo}: {e}")

    await asyncio.get_event_loop().run_in_executor(None, _do_cleanup)


# ─── Endpoints ────────────────────────────────────────────

@router.post("/upload-documento/")
@limiter.limit("10/minute")
async def upload_documento(
    request: Request,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    current_user: User = Depends(get_optional_current_user),
):
    background_tasks.add_task(limpiar_archivos_antiguos)

    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Solo se aceptan archivos .docx")

    contents   = await file.read()
    upload_id  = str(uuid.uuid4())
    safe_name  = RE_SAFE_FILENAME.sub("_", file.filename) if file.filename else "upload.docx"
    input_path = os.path.join(UPLOAD_DIR, f"{upload_id}_{safe_name}")

    with open(input_path, "wb") as f:
        f.write(contents)

    upload_storage.set(upload_id, (input_path, safe_name))
    logger.info(f"Upload #{upload_id}: {safe_name} ({len(contents)} bytes)")
    return {"upload_id": upload_id, "filename": safe_name}


@router.get("/procesar-apa/stream")
@limiter.limit("10/minute")
async def procesar_apa_stream(
    request: Request,
    upload_id: str = Query(...),
    edicion: str   = Query("7ma"),
    plan: str      = Query("free"),
    token: str     = Query(None),
    db: Session    = Depends(get_db),
):
    current_user = None
    if token and token != "null":
        current_user = _decode_user_from_token(token, db)

    entry = upload_storage.get(upload_id)
    if entry is None:
        raise HTTPException(status_code=404, detail="upload_id no encontrado. Sube el archivo primero.")
    input_path, filename = entry

    if plan == "pro":
        if not current_user:
            raise HTTPException(status_code=401, detail="Debes iniciar sesión para usar DocAI Pro.")
        if get_available_tokens(current_user.id, db)["total"] <= 0:
            raise HTTPException(status_code=402, detail="Sin tokens disponibles.")

    try:
        with open(input_path, "rb") as f:
            doc = Document(io.BytesIO(f.read()))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"No se pudo leer el .docx: {e}")

    async def event_generator() -> AsyncGenerator[str, None]:
        try:
            if plan == "pro":
                async for evento in procesar_con_ia_stream(doc.paragraphs):
                    if evento.get("tipo") == "finalizado":
                        consume_tokens(current_user.id, evento.get("groq_tokens", 0), filename, db)
                        # No eliminamos input_path aquí: /generar-final/ lo necesita
                        # para copiar la portada con imágenes. El cron de limpieza
                        # (limpiar_archivos_antiguos) lo borrará después de 24h.
                        try:
                            upload_storage.pop(upload_id)
                        except Exception:
                            pass
                    yield f"data: {json.dumps(evento, ensure_ascii=False)}\n\n"
            else:
                resultado = procesar_con_reglas(doc.paragraphs)
                yield f"data: {json.dumps({'tipo': 'inicio', 'total_lotes': 1, 'progreso': 0, 'modelo': 'reglas'}, ensure_ascii=False)}\n\n"
                yield f"data: {json.dumps({'tipo': 'finalizado', 'progreso': 100, 'stats': resultado['stats'], 'detalles': resultado['detalles'], 'groq_tokens': 0}, ensure_ascii=False)}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'tipo': 'error', 'mensaje': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@router.post("/procesar-apa/")
async def procesar_documento(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    edicion: str = Form("7ma"),
    plan: str = Form("free"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_optional_current_user),
):
    background_tasks.add_task(limpiar_archivos_antiguos)

    contents   = await file.read()
    safe_name  = RE_SAFE_FILENAME.sub("_", file.filename) if file.filename else "upload.docx"
    input_path = os.path.join(UPLOAD_DIR, f"{uuid.uuid4().hex}_{safe_name}")

    with open(input_path, "wb") as f:
        f.write(contents)

    if not os.path.exists(input_path):
        raise HTTPException(status_code=500, detail="Error al guardar el archivo")

    try:
        doc = Document(io.BytesIO(contents))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"No se pudo procesar el .docx: {e}")

    logger.info(f"Procesando: {safe_name} (Plan: {plan}, Edicion: {edicion})")

    if plan == "pro":
        balance = get_available_tokens(current_user.id, db)
        if balance["total"] <= 0:
            raise HTTPException(status_code=402, detail="No tienes tokens disponibles.")
        resultado   = procesar_con_ia(doc.paragraphs)
        groq_tokens = resultado.get('groq_tokens', 0)
        consume_tokens(current_user.id, groq_tokens, safe_name, db)
    else:
        resultado = procesar_con_reglas(doc.paragraphs)
        resultado["groq_tokens"] = 0

    return {
        "status": "success",
        "plan": plan,
        "resumen": resultado["stats"],
        "detalles": resultado["detalles"],
        "tokens_consumed": groq_tokens_to_docai(resultado.get("groq_tokens", 0)),
    }


@router.get("/tokens/balance")
async def mis_tokens(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    return {"status": "success", **get_available_tokens(current_user.id, db)}


@router.post("/generar-final/")
async def generar_final(
    datos: DatosFinales,
    current_user: User = Depends(get_current_user),
):
    base_name     = os.path.splitext(datos.filename)[0]
    safe_base     = RE_SAFE_BASENAME.sub("_", base_name).strip()
    unique_suffix = uuid.uuid4().hex
    out_name      = f"FINAL_{datos.edicion}_{safe_base}_{unique_suffix}"
    out_docx      = os.path.join(PROCESSED_DIR, out_name + ".docx")
    output_path   = out_docx

    # ── Recuperar documento original (para portada con imágenes) ──────
    doc_original = None
    if datos.upload_id:
        orig_path = None

        # 1. Buscar en el storage en memoria (puede haberse expirado/popado)
        entry = upload_storage.get(datos.upload_id)
        if entry:
            orig_path = entry[0]

        # 2. Fallback: buscar el archivo en disco por prefijo de upload_id
        if not orig_path or not os.path.exists(orig_path):
            prefix = datos.upload_id
            try:
                for fname in os.listdir(UPLOAD_DIR):
                    if fname.startswith(prefix):
                        orig_path = os.path.join(UPLOAD_DIR, fname)
                        break
            except Exception:
                pass

        if orig_path and os.path.exists(orig_path):
            try:
                with open(orig_path, "rb") as f:
                    doc_original = Document(io.BytesIO(f.read()))
            except Exception as e:
                logger.warning(f"No se pudo abrir el original para copiar portada: {e}")

    # ── Detectar cuántos párrafos son portada ─────────────────────────
    # Primero intentamos con el valor enviado por el frontend; si es 0
    # lo calculamos automáticamente desde la lista de párrafos.
    n_portada = datos.n_portada
    if n_portada == 0 and doc_original:
        n_portada = _detectar_n_portada(
            [p.dict() for p in datos.parrafos]
        )

    doc    = Document()
    reglas = dict(NORMAS_APA.get(datos.edicion, NORMAS_APA["7ma"]))
    reglas["edicion"] = datos.edicion
    if datos.edicion == "6ta":
        reglas["fuente"] = DEFAULT_APA_FONT
        reglas["tamano"] = FUENTES_APA[DEFAULT_APA_FONT]["tamano"]
    else:
        reglas["fuente"] = validar_fuente_apa(datos.fuente)
        reglas["tamano"] = FUENTES_APA[reglas["fuente"]]["tamano"]

    for section in doc.sections:
        section.page_width = section.page_height = None
        section.page_width  = LETTER_PAGE["width"]
        section.page_height = LETTER_PAGE["height"]
        m = Inches(reglas["margen"])
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = m

    ns = doc.styles["Normal"]
    ns.font.name = reglas["fuente"]
    ns.font.size = Pt(reglas["tamano"])
    nspf = ns.paragraph_format
    nspf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    nspf.space_before = nspf.space_after = Pt(0)
    nspf.first_line_indent = Inches(reglas["sangria_primera_linea"])
    nspf.left_indent = Inches(0)

    _configurar_encabezado_paginas(doc)

    # ═══════════════════════════════════════════════════════
    # ORDEN APA: 1) Portada  2) Índice  3) Cuerpo
    # ═══════════════════════════════════════════════════════

    # --- 1. Portada (copiada del original con imágenes) ---
    if doc_original and n_portada > 0:
        try:
            copiar_portada_desde_original(doc_original, doc, n_portada)
            doc.add_page_break()
            logger.info(f"Portada copiada: {n_portada} párrafos del original")
        except Exception as e:
            logger.warning(f"Error copiando portada, se omite: {e}")

    # --- 2. Índice (solo Pro) ---
    if datos.incluir_indice and datos.plan == "pro":
        headings = []
        for p in datos.parrafos[n_portada:]:
            cat = _normalizar_categoria(p.categoria)
            if "TITULO" in cat:
                try:
                    nivel = int(cat.split("_")[-1].replace("N", ""))
                except Exception:
                    nivel = 1
                headings.append((min(max(nivel, 1), 3), p.texto.strip()))

        title = doc.add_paragraph("Índice")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.bold, run.font.name, run.font.size = True, reglas["fuente"], Pt(16)

        if headings:
            _insertar_tabla_de_contenidos(doc)
            nota = doc.add_paragraph("Actualiza los campos en Word para obtener los números de página reales.")
            nota.italic = True
            nota.paragraph_format.space_before = Pt(4)
            nota.paragraph_format.space_after  = Pt(8)
        else:
            doc.add_paragraph("No se detectaron títulos válidos.").italic = True

        doc.add_page_break()

    # --- 3. Cuerpo del documento (párrafos APA, sin los de portada) ---
    reference_started = False
    paragraph_counter = 0
    reference_buffer: list[ParrafoCorregido] = []
    i = n_portada          # ← Empezamos DESPUÉS de los párrafos de portada
    parrafos = datos.parrafos

    while i < len(parrafos):
        p   = parrafos[i]
        cat = _normalizar_categoria(p.categoria)

        # Títulos N3-N5 en 6ta: fusionar con párrafo siguiente
        if datos.edicion == "6ta" and cat in {"TITULO_N3", "TITULO_N4", "TITULO_N5"} and i + 1 < len(parrafos):
            if _normalizar_categoria(parrafos[i + 1].categoria) == "PARRAFO_NORMAL":
                paragraph = doc.add_paragraph()
                configurar_parrafo_estilo(paragraph, cat, reglas, body_text=parrafos[i + 1].texto.strip())
                paragraph_counter += 1
                i += 2
                continue

        # Detectar encabezado de referencias
        if not reference_started and _es_encabezado_referencias(p.texto):
            if paragraph_counter > 0:
                doc.add_page_break()
            heading_texto = p.texto.strip()
            if i + 1 < len(parrafos) and _es_continuacion_encabezado_referencias(parrafos[i + 1].texto):
                heading_texto += " " + parrafos[i + 1].texto.strip()
                i += 1
            ref_h = doc.add_paragraph(heading_texto)
            ref_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in ref_h.runs:
                run.bold = datos.edicion != "6ta"
                run.font.name, run.font.size = reglas["fuente"], Pt(reglas["tamano"])
            ref_h.paragraph_format.space_before = Pt(12)
            ref_h.paragraph_format.space_after  = Pt(12)
            reference_started = True
            paragraph_counter += 1
            i += 1
            continue

        if cat == "REFERENCIA" and not reference_started:
            if paragraph_counter > 0:
                doc.add_page_break()
            ref_h = doc.add_paragraph("Referencias")
            ref_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in ref_h.runs:
                run.bold = datos.edicion != "6ta"
                run.font.name, run.font.size = reglas["fuente"], Pt(reglas["tamano"])
            ref_h.paragraph_format.space_before = Pt(12)
            ref_h.paragraph_format.space_after  = Pt(12)
            reference_started = True

        if reference_started and cat == "REFERENCIA":
            reference_buffer.append(p)
            i += 1
            continue

        if reference_started and reference_buffer:
            for ref in sorted(reference_buffer, key=lambda r: _ordenar_referencia_por_autor(r.texto)):
                ph = doc.add_paragraph(ref.texto)
                configurar_parrafo_estilo(ph, ref.categoria, reglas)
                paragraph_counter += 1
            reference_buffer = []

        ph = doc.add_paragraph(p.texto)
        configurar_parrafo_estilo(ph, cat, reglas)
        # Aplicar alineación personalizada del usuario (si la especificó en el editor)
        if p.textAlign:
            _MAP_ALIGN = {
                'left':   WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right':  WD_ALIGN_PARAGRAPH.RIGHT,
            }
            if p.textAlign in _MAP_ALIGN:
                ph.alignment = _MAP_ALIGN[p.textAlign]
        paragraph_counter += 1
        i += 1

    if reference_buffer:
        for ref in sorted(reference_buffer, key=lambda r: _ordenar_referencia_por_autor(r.texto)):
            ph = doc.add_paragraph(ref.texto)
            configurar_parrafo_estilo(ph, ref.categoria, reglas)
            paragraph_counter += 1

    if datos.incluir_indice and datos.plan == "pro":
        _force_update_fields(doc)
    if datos.plan == "free":
        añadir_marca_de_agua(doc)

    try:
        doc.save(out_docx)
    except PermissionError:
        raise HTTPException(status_code=500, detail="No se pudo guardar el archivo DOCX.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"No se pudo generar el archivo DOCX: {e}")

    # --- Conversión a PDF (solo Pro) ---
    if datos.formato.lower() == "pdf":
        if datos.plan != "pro":
            raise HTTPException(status_code=403, detail="PDF exclusivo para usuarios Pro.")

        out_pdf = os.path.abspath(os.path.join(PROCESSED_DIR, out_name + ".pdf"))
        if os.path.exists(out_pdf):
            try:
                os.remove(out_pdf)
            except Exception:
                pass

        soffice = _get_soffice_path_cached()
        if not soffice:
            raise HTTPException(status_code=500, detail="LibreOffice no está instalado en el servidor.")

        try:
            tmp_dir = tempfile.mkdtemp()
            result  = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", tmp_dir,
                 os.path.abspath(out_docx)],
                capture_output=True, text=True, timeout=60,
            )
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice falló: {result.stderr}")
            generated = os.path.join(tmp_dir, os.path.splitext(os.path.basename(out_docx))[0] + ".pdf")
            if not os.path.exists(generated):
                raise RuntimeError("LibreOffice no generó el PDF esperado.")
            shutil.move(generated, out_pdf)
            shutil.rmtree(tmp_dir, ignore_errors=True)
            output_path = out_pdf
        except subprocess.TimeoutExpired:
            raise HTTPException(status_code=500, detail="La conversión a PDF tardó demasiado.")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"No se pudo convertir a PDF: {e}")

    file_id = str(uuid.uuid4())
    storage.set(file_id, output_path)
    return {"file_id": file_id}


@router.get("/imagen/{upload_id}/{rel_id}")
async def obtener_imagen_portada(upload_id: str, rel_id: str):
    """
    Sirve una imagen extraída del .docx original identificado por upload_id.
    Se usa para mostrar imágenes de portada en la vista previa del editor.
    """
    from fastapi.responses import Response
    import base64

    # Buscar el archivo en upload_storage o en disco
    orig_path = None
    entry = upload_storage.get(upload_id)
    if entry:
        orig_path = entry[0]

    if not orig_path or not os.path.exists(orig_path):
        try:
            for fname in os.listdir(UPLOAD_DIR):
                if fname.startswith(upload_id):
                    orig_path = os.path.join(UPLOAD_DIR, fname)
                    break
        except Exception:
            pass

    if not orig_path or not os.path.exists(orig_path):
        raise HTTPException(status_code=404, detail="Archivo original no encontrado")

    try:
        with open(orig_path, "rb") as f:
            doc = Document(io.BytesIO(f.read()))

        img_part = doc.part.related_parts.get(rel_id)
        if img_part is None:
            raise HTTPException(status_code=404, detail="Imagen no encontrada")

        return Response(
            content=img_part._blob,
            media_type=img_part.content_type,
            headers={"Cache-Control": "public, max-age=3600"}
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extrayendo imagen: {e}")


@router.get("/descargar/{file_id}")
async def descargar_archivo(file_id: str):
    path = storage.get(file_id)
    if path and os.path.exists(path):
        return FileResponse(path=path, filename=os.path.basename(path))
    raise HTTPException(status_code=404, detail="No encontrado")

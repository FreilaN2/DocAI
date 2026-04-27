from docx import Document
from docx.shared import Inches

def generar_documento_extenso(nombre_archivo):
    doc = Document()
    
    # --- PÁGINA 1: Portada y Título Principal ---
    # Gemini debería detectarlo como TITULO_N1
    doc.add_paragraph("Impacto de la Inteligencia Artificial en la Gestión de Proyectos de SYSCOTEK")
    
    # Texto de relleno para simular espacio de portada
    for _ in range(5): doc.add_paragraph("") 

    # --- PÁGINA 1-2: Introducción y Marco Teórico ---
    doc.add_paragraph("Introducción") # TITULO_N2
    
    texto_largo = (
        "La implementación de sistemas inteligentes en el entorno empresarial de SYSCOTEK "
        "representa un cambio de paradigma en la eficiencia operativa. En la actualidad, "
        "la gestión de proyectos requiere una integración fluida entre el backend y las "
        "necesidades del cliente final. El uso de modelos como Gemini 3.1 permite que "
        "tareas que antes tomaban horas, como el formateo de documentos académicos, "
        "se realicen en segundos con una precisión semántica sin precedentes. "
    )
    # Repetimos para generar volumen de texto
    for _ in range(4): doc.add_paragraph(texto_largo * 2)

    doc.add_paragraph("Antecedentes de la Automatización") # TITULO_N2
    doc.add_paragraph(
        "Históricamente, la automatización se limitaba a scripts rígidos. Sin embargo, "
        "con la llegada de la IA generativa, podemos interpretar el contexto del usuario. "
        "En SYSCOTEK, hemos observado que la integración de Python con sistemas ERP "
        "como Odoo facilita la recolección de datos en tiempo real."
    )
    for _ in range(3): doc.add_paragraph(texto_largo)

    # --- PÁGINA 2: Metodología y Niveles Profundos ---
    doc.add_paragraph("Metodología de la Investigación") # TITULO_N2
    doc.add_paragraph(
        "Se procedió a realizar un análisis exhaustivo de las herramientas disponibles "
        "en el mercado actual para la automatización de procesos administrativos."
    )

    doc.add_paragraph("Fase de Recolección de Datos") # TITULO_N3
    doc.add_paragraph(
        "Durante esta etapa, se utilizaron APIs de Google para extraer información "
        "relevante de documentos no estructurados. Se aplicaron técnicas de Prompt Engineering "
        "para asegurar que el modelo clasificara correctamente cada párrafo."
    )

    doc.add_paragraph("Análisis de Integración con Odoo") # TITULO_N3
    doc.add_paragraph(
        "La conexión entre el módulo de proyectos de Odoo y el backend de FastAPI "
        "permitió una trazabilidad completa de los cambios realizados por los usuarios."
    )

    # --- PÁGINA 3: Conclusiones y Referencias ---
    doc.add_paragraph("Resultados y Discusión") # TITULO_N2
    doc.add_paragraph(
        "Los resultados indican que DocAI reduce el tiempo de corrección de tesis en un 80%. "
        "Los usuarios prefieren la validación manual a través de interfaces en React "
        "antes de la generación definitiva del archivo Word."
    )
    for _ in range(4): doc.add_paragraph(texto_largo)

    doc.add_paragraph("Referencias Bibliográficas") # TITULO_N2
    
    # Varias referencias para probar la Sangría Francesa
    referencias = [
        "Cote, M. (2026). Guía avanzada de integración de Odoo y Python. Editorial Tech.",
        "García, J. & Pérez, L. (2025). Inteligencia Artificial aplicada a la empresa. McGraw-Hill.",
        "SYSCOTEK Labs. (2026). Reporte anual de innovación tecnológica y automatización.",
        "Wilson, K. (2024). The Future of Academic Formatting and AI. Journal of Tech."
    ]
    for ref in referencias:
        doc.add_paragraph(ref)

    # Forzamos márgenes incorrectos (0.5") para que DocAI los lleve a 1.0"
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    doc.save(nombre_archivo)
    print(f"✅ Archivo extenso '{nombre_archivo}' generado (~3 páginas).")
    print("🚀 Súbelo para probar el visor de scroll y los niveles de títulos.")

generar_documento_extenso("test_extenso_docai.docx")